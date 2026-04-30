#!/usr/bin/env python3
"""
Legacy internal DOCX generator for Google_Ads_Agent review notes.

Do not use this for branded client-facing rebuild reviews. The canonical
client deliverable is HTML designed first, then exported to PDF through
shared/presentation/build_review_doc.py. See docs/HTML_PDF_CLIENT_REPORT_STANDARD.md.
"""

from __future__ import annotations

import argparse
import csv
import datetime as dt
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]

BRAND = {
    "teal": "A9CBD1",
    "sage": "7E938D",
    "peach": "E9BCA8",
    "cream": "F7F5EF",
    "ink": "384447",
    "muted": "6E7D7B",
    "line": "DCE5E2",
    "white": "FFFFFF",
}


def read_editor_csv(path: Path) -> tuple[list[dict[str, str]], list[str]]:
    with path.open("r", encoding="utf-16", newline="") as f:
        reader = csv.DictReader(f, dialect=csv.excel_tab)
        rows = list(reader)
        return rows, list(reader.fieldnames or [])


def truthy(value: str | None) -> bool:
    return bool((value or "").strip())


def add_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    size: float = 9,
    color: str = BRAND["ink"],
    margin: int = 120,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_cell_margins(cell, margin, margin, margin, margin)


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = RGBColor.from_string(BRAND["ink"])
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 24, BRAND["ink"]),
        ("Heading 1", 15, BRAND["sage"]),
        ("Heading 2", 11, BRAND["ink"]),
        ("Heading 3", 10, BRAND["sage"]),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_band(doc: Document, title: str, subtitle: str = "") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    add_shading(cell, BRAND["cream"])
    add_cell_margins(cell, 180, 180, 180, 180)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.name = "Aptos Display"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BRAND["ink"])
    if subtitle:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(subtitle)
        r2.font.name = "Aptos"
        r2.font.size = Pt(8.8)
        r2.font.color.rgb = RGBColor.from_string(BRAND["muted"])


def add_metric_cards(doc: Document, metrics: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(metrics))
    table.autofit = True
    for idx, (value, label, note) in enumerate(metrics):
        cell = table.cell(0, idx)
        add_shading(cell, BRAND["cream"] if idx % 2 == 0 else "EEF5F4")
        add_cell_margins(cell, 160, 140, 160, 140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        r.font.name = "Aptos Display"
        r.font.size = Pt(17)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(BRAND["sage"])
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.name = "Aptos"
        r2.font.size = Pt(8.5)
        r2.font.bold = True
        r2.font.color.rgb = RGBColor.from_string(BRAND["ink"])
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(note)
        r3.font.name = "Aptos"
        r3.font.size = Pt(7.5)
        r3.font.color.rgb = RGBColor.from_string(BRAND["muted"])


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
    font_size: float = 8,
    margin: int = 120,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        add_shading(hdr[idx], BRAND["teal"])
        set_cell_text(hdr[idx], h, bold=True, size=font_size + 0.2, color=BRAND["ink"], margin=margin)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=font_size, margin=margin)
            if idx == 0:
                add_shading(cells[idx], "F7FAFA")
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                set_cell_width(row.cells[idx], width)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.18)
    r = p.add_run("- ")
    r.font.color.rgb = RGBColor.from_string(BRAND["peach"])
    r.font.bold = True
    t = p.add_run(text)
    t.font.name = "Aptos"
    t.font.size = Pt(9)
    t.font.color.rgb = RGBColor.from_string(BRAND["ink"])


def page_break(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_cover(doc: Document, client: str, website: str, agency: str, screenshot: Path | None) -> None:
    if screenshot and screenshot.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(screenshot), width=Inches(6.9))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Campaign Rebuild Review")
    r.font.name = "Aptos Display"
    r.font.size = Pt(25)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BRAND["ink"])
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run(client)
    s.font.name = "Aptos Display"
    s.font.size = Pt(15)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(BRAND["sage"])
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(f"{website} | Prepared by {agency} | {dt.date.today().isoformat()}")
    m.font.name = "Aptos"
    m.font.size = Pt(9)
    m.font.color.rgb = RGBColor.from_string(BRAND["muted"])
    add_band(
        doc,
        "Human Review Draft",
        "This document summarizes the rebuild logic, sample ads, and approval points before Google Ads Editor posting.",
    )


def collect_data(current_rows: list[dict[str, str]], old_rows: list[dict[str, str]] | None) -> dict:
    ad_groups = sorted({r.get("Ad Group", "") for r in current_rows if truthy(r.get("Ad Group"))})
    keyword_rows = [r for r in current_rows if truthy(r.get("Keyword"))]
    rsa_rows = [r for r in current_rows if r.get("Ad type") == "Responsive search ad"]
    location_rows = [r for r in current_rows if truthy(r.get("Location"))]
    service_counts = Counter([a.rsplit(" - ", 1)[0] for a in ad_groups if " - " in a])

    old = {}
    if old_rows:
        old = {
            "campaigns": len({r.get("Campaign") for r in old_rows if truthy(r.get("Campaign"))}),
            "ad_groups": len({r.get("Ad Group") for r in old_rows if truthy(r.get("Ad Group"))}),
            "keywords": sum(1 for r in old_rows if truthy(r.get("Keyword"))),
            "rsas": sum(1 for r in old_rows if r.get("Ad type") == "Responsive search ad"),
        }

    return {
        "ad_groups": ad_groups,
        "keyword_rows": keyword_rows,
        "rsa_rows": rsa_rows,
        "location_rows": location_rows,
        "service_counts": service_counts,
        "old": old,
    }


def ad_group_examples(data: dict, preferred: list[str]) -> list[dict[str, str]]:
    rows = data["rsa_rows"]
    by_ag = {r.get("Ad Group", ""): r for r in rows}
    selected = []
    for ag in preferred:
        if ag in by_ag:
            selected.append(by_ag[ag])
    if len(selected) < 4:
        for r in rows:
            if r not in selected:
                selected.append(r)
            if len(selected) >= 4:
                break
    return selected[:4]


def headlines(row: dict[str, str]) -> list[str]:
    return [row.get(f"Headline {i}", "") for i in range(1, 16) if truthy(row.get(f"Headline {i}", ""))]


def descriptions(row: dict[str, str]) -> list[str]:
    return [row.get(f"Description {i}", "") for i in range(1, 5) if truthy(row.get(f"Description {i}", ""))]


def add_ad_preview(doc: Document, row: dict[str, str]) -> None:
    ag = row.get("Ad Group", "Ad Group")
    final_url = row.get("Final URL", "")
    h = headlines(row)
    d = descriptions(row)[:2]
    add_band(doc, ag, f"Landing page: {final_url}")
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    add_shading(cell, "FFFFFF")
    add_cell_margins(cell, 150, 170, 150, 170)
    p = cell.paragraphs[0]
    r0 = p.add_run("Sponsored  ")
    r0.font.name = "Aptos"
    r0.font.size = Pt(7.5)
    r0.font.color.rgb = RGBColor.from_string(BRAND["muted"])
    url = p.add_run(final_url.replace("https://", "").split("/")[0])
    url.font.name = "Aptos"
    url.font.size = Pt(7.5)
    url.font.color.rgb = RGBColor.from_string(BRAND["muted"])
    p2 = cell.add_paragraph()
    title = " | ".join(h[:3])
    rt = p2.add_run(title)
    rt.font.name = "Aptos Display"
    rt.font.size = Pt(11.5)
    rt.font.bold = True
    rt.font.color.rgb = RGBColor(26, 88, 150)
    p3 = cell.add_paragraph()
    rd = p3.add_run(" ".join(d))
    rd.font.name = "Aptos"
    rd.font.size = Pt(8.6)
    rd.font.color.rgb = RGBColor.from_string(BRAND["ink"])
    add_table(
        doc,
        ["#", "Headlines for review", "Characters"],
        [[f"{idx + 1}", x, str(len(x))] for idx, x in enumerate(h)],
        widths=[0.45, 5.0, 0.9],
        font_size=7.25,
        margin=60,
    )
    add_table(
        doc,
        ["Descriptions for review", "Characters"],
        [[x, str(len(x))] for x in descriptions(row)],
        widths=[5.4, 1.0],
        font_size=7.25,
        margin=60,
    )


def build_doc(args: argparse.Namespace) -> Path:
    current_rows, _ = read_editor_csv(Path(args.current_csv))
    old_rows = None
    if args.previous_export:
        old_rows, _ = read_editor_csv(Path(args.previous_export))
    data = collect_data(current_rows, old_rows)

    doc = Document()
    style_doc(doc)
    add_cover(doc, args.client_name, args.website, args.agency_name, Path(args.screenshot) if args.screenshot else None)

    page_break(doc)
    doc.add_heading("Executive Summary", level=1)
    metrics = [
        (str(len(data["service_counts"])), "Service categories", "Structured into reviewable service families"),
        (str(len(data["ad_groups"])), "Ad groups", "Organized by service and intent layer"),
        (str(len(data["keyword_rows"])), "Phrase keywords", "No broad match in the rebuild"),
        (str(len(data["rsa_rows"])), "Responsive ads", "One RSA per ad group for review"),
    ]
    add_metric_cards(doc, metrics)
    bullet(doc, "The rebuild turns the account into a clean Search campaign structure that can be staged in Google Ads Editor.")
    bullet(doc, "Ad groups are organized by service, audience or category, and search intent layer.")
    bullet(doc, "Keywords are phrase match only, with city, local, state, and general variants separated for cleaner reporting.")
    bullet(doc, "Ad copy uses complete RSA assets with stronger headlines, service relevance, and client review points.")

    doc.add_heading("Before And After", level=1)
    old = data["old"]
    before_counts = "Existing export" if not old else f"{old['campaigns']} campaigns, {old['ad_groups']} ad groups, {old['keywords']} keyword rows"
    add_table(
        doc,
        ["Starting Point", "Rebuild Approach", "Why It Matters"],
        [
            [before_counts, "One staged Search rebuild organized around service and intent.", "Cleaner approval, cleaner reporting, and safer launch control."],
            ["Mixed search intent inside broader groups.", "General, Local, City, and State ad groups are separated.", "We can see which intent layer actually performs."],
            ["Search terms were useful but not systematically mapped.", "Converting and relevant terms inform phrase keyword expansion.", "The build learns from account history without blindly copying noise."],
            ["Ad copy quality varied by ad group.", "RSA copy uses headline and description gates before staging.", "The client reviews stronger examples before upload."],
        ],
        widths=[2.05, 2.35, 2.2],
    )
    if args.previous_provider:
        bullet(
            doc,
            f"The imported export included {args.previous_provider} in legacy naming. We treat that as source attribution only and remove it from the rebuilt campaign naming system.",
        )

    doc.add_heading("Campaign Architecture", level=1)
    add_table(
        doc,
        ["Layer", "Pattern", "Purpose", "Example"],
        [
            ["General", "Service - Category - General", "Core non-geo searches.", "Psychiatry - Adult - General"],
            ["Local", "Service - Category - Local", "Near-me and close-proximity searches.", "Therapy - Trauma - Local"],
            ["City", "Service - Category - City", "Named city intent across approved target cities.", "Testing - ADHD - City"],
            ["State", "Service - Category - State", "State and regional modifiers.", "Psychiatry - Children - State"],
        ],
        widths=[0.9, 1.8, 2.2, 1.7],
    )
    services = ", ".join(list(data["service_counts"].keys())[:12])
    bullet(doc, f"Representative service families include: {services}.")
    bullet(doc, "The first build keeps service coverage together for staging, then future data can justify splitting by service line or geography.")

    page_break(doc)
    doc.add_heading("Keyword And Targeting Strategy", level=1)
    add_table(
        doc,
        ["Decision", "Current Build", "Client Review Point"],
        [
            ["Match type", "Phrase match only.", "Confirm this conservative launch control."],
            ["City coverage", "Every city ad group receives all approved city modifiers.", "Confirm target cities and any exclusions."],
            ["Local intent", "Near-me variations are isolated in Local ad groups.", "Confirm whether local-intent traffic should receive budget priority."],
            ["Locations", f"{len(data['location_rows'])} location or radius rows in the staging CSV.", "Confirm core markets and bid adjustments."],
        ],
        widths=[1.5, 2.6, 2.4],
    )
    sample_keywords = []
    for ag in ["Psychiatry - Adult - General", "Psychiatry - Adult - City", "Therapy - Trauma - Local", "Testing - ADHD - General"]:
        kws = [r.get("Keyword", "") for r in data["keyword_rows"] if r.get("Ad Group") == ag][:5]
        if kws:
            sample_keywords.append([ag, ", ".join(kws)])
    add_table(doc, ["Sample Ad Group", "Example Keywords"], sample_keywords, widths=[2.2, 4.3])

    doc.add_heading("How We Split Over Time", level=1)
    add_table(
        doc,
        ["Split Trigger", "What We Would Do", "Reason"],
        [
            ["Service volume", "Move a winning service family into its own campaign.", "Separate budget and bidding control."],
            ["Geo performance", "Split high-performing cities or counties.", "Protect efficient markets and reduce wasted spend."],
            ["Near-me volume", "Build a local-intent campaign.", "Give near-me searches their own budget and message testing."],
            ["ZIP evidence", "Move from city/county targeting into ZIP-level structure.", "Only when enough data supports the extra complexity."],
        ],
        widths=[1.7, 2.6, 2.2],
    )

    page_break(doc)
    doc.add_heading("Ad Copy Review", level=1)
    bullet(doc, "The examples below show actual RSA assets from the staging build.")
    bullet(doc, "The client should review service fit, claims, tone, landing pages, and any phrases that need adjustment.")
    examples = ad_group_examples(
        data,
        [
            "Psychiatry - Adult - General",
            "Psychiatry - Adult - City",
            "Therapy - Trauma - City",
            "Testing - ADHD - General",
        ],
    )
    for idx, row in enumerate(examples):
        if idx:
            page_break(doc)
        add_ad_preview(doc, row)

    page_break(doc)
    doc.add_heading("Approval Checklist", level=1)
    add_table(
        doc,
        ["Review Area", "Approve", "Needs Changes", "Notes"],
        [
            ["Service categories and ad group organization", "[ ]", "[ ]", ""],
            ["Target cities, counties, states, and radius targets", "[ ]", "[ ]", ""],
            ["Landing page assignments", "[ ]", "[ ]", ""],
            ["Ad headlines and descriptions", "[ ]", "[ ]", ""],
            ["Claims, credentials, offers, and wording", "[ ]", "[ ]", ""],
            ["Ready to import into Google Ads Editor", "[ ]", "[ ]", ""],
        ],
        widths=[3.0, 0.8, 1.2, 1.8],
    )
    doc.add_heading("Sources And Staging Notes", level=1)
    bullet(doc, "Google Ads Editor is the staging environment. The CSV should be reviewed in Editor before posting.")
    bullet(doc, "Google Ads Editor CSV import: https://support.google.com/google-ads/editor/answer/30564?hl=en")
    bullet(doc, "Google Ads Editor CSV columns: https://support.google.com/google-ads/editor/answer/57747?hl=en")
    bullet(doc, "Responsive search ads: https://support.google.com/google-ads/answer/7684791?hl=en-EN")
    bullet(doc, "Google geo targets: https://developers.google.com/google-ads/api/data/geotargets")

    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate a legacy internal rebuild review DOCX")
    parser.add_argument("--current-csv", required=True)
    parser.add_argument("--previous-export")
    parser.add_argument("--client-name", required=True)
    parser.add_argument("--website", required=True)
    parser.add_argument("--agency-name", default="Advertising Report Card")
    parser.add_argument("--previous-provider")
    parser.add_argument("--screenshot")
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    out = build_doc(args)
    print(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
